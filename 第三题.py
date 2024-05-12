import json
import requests
from docx import Document
from docx.shared import Pt

json_file_path = r""
template_path = r""
output_path = r""

file = open(json_file_path, 'r', encoding='utf-8')
data = json.load(file)
file.close()

api_url = ""
headers = {'Content-Type': 'application/json'}
response = requests.post(api_url, json=data, headers=headers)

doc = Document(template_path)
for paragraph in doc.paragraphs:
    if '公司名称' in paragraph.text:
        paragraph.text = paragraph.text.replace('公司名称', data['campany'])
    if '公司法人' in paragraph.text:
        paragraph.text = paragraph.text.replace('公司法人', data['legalName'])
    if '法人联系方式' in paragraph.text:
        paragraph.text = paragraph.text.replace('法人联系方式', data['legalPhone'])

employee_table = doc.tables[-1]
employee_list = data.get('employeeInfor', [])
for i in range(len(employee_list)):
    employee = employee_list[i]
    row = employee_table.rows[i]
    cells = row.cells
    cells[0].text = f"员工{i+1}"
    cells[1].text = employee['name']
    cells[2].text = employee['idcard']
    cells[3].text = employee['birthday']

for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith('Heading'):
        paragraph.alignment = 0  # 设置为左对齐
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
    if paragraph == doc.paragraphs[0]:
        paragraph.style = 'Heading1'
        paragraph.alignment = 1  # 设置为居中对齐

doc.save(output_path)